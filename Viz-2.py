!pip install python-docx
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches

def plot_temperature_data(csv_file, date_column, temperature_columns, city_names):
    # Read the temperature data from CSV
    temperature_data = pd.read_csv(csv_file)
    
    # Convert date column to datetime format
    temperature_data[date_column] = pd.to_datetime(temperature_data[date_column])

    # Plotting
    plt.figure(figsize=(10, 6))

    # Plot temperature data for each city
    for temp_col, city_name in zip(temperature_columns, city_names):
        plt.plot(temperature_data[date_column], temperature_data[temp_col], label=city_name)

    # Customize the plot
    plt.xlabel('Date')
    plt.ylabel('Temperature (°C)')
    plt.title('Avg monthly temparature of popular 4 cities over last 2 years')
    plt.legend()
    plt.grid(True)
    plt.tight_layout()

    # Show the plot
    return plt

def plot_stacked_column_chart(data_file):
    # Read the dataset
    df = pd.read_csv(data_file)

    # Filter data from the year 2000 onwards
    df = df[df['release_year'] >= 2000]

    # Extract release years, count movies and TV shows for each year
    grouped_data = df.groupby(['release_year', 'type']).size().unstack().fillna(0)

    # Plotting
    plt.figure(figsize=(10, 6))
    plt.bar(grouped_data.index, grouped_data['Movie'], color='#FFA500', label='Movies')
    plt.bar(grouped_data.index, grouped_data['TV Show'], color='#32CD32', label='TV Shows', bottom=grouped_data['Movie'])

    # Customize the plot
    plt.xlabel('Year')
    plt.ylabel('Number of Titles')
    plt.title('Number of Movies and TV Shows Released Each Year (Starting from 2000)')
    plt.legend()
    plt.grid(True)
    plt.tight_layout()

    # Show the plot
    return plt

def plot_gender_usage_pie_chart(data_file):
    # Read the dataset
    df = pd.read_csv(data_file)

    # Filter data for the years 2019 and 2020
    df_2019 = df[df['Year'] == 2019]
    df_2020 = df[df['Year'] == 2020]

    # Count the number of males and females for each year
    male_count_2019 = df_2019['Gender_Coded'].value_counts().get(1, 0)
    female_count_2019 = df_2019['Gender_Coded'].value_counts().get(2, 0)

    male_count_2020 = df_2020['Gender_Coded'].value_counts().get(1, 0)
    female_count_2020 = df_2020['Gender_Coded'].value_counts().get(2, 0)

    # Data for pie charts
    male_counts = [male_count_2019, male_count_2020]
    female_counts = [female_count_2019, female_count_2020]
    years = ['2019', '2020']

    # Plotting side-by-side pie charts
    fig, axes = plt.subplots(1, 2, figsize=(10, 5))

    # Pie chart for Male vs Female in 2019
    axes[0].pie(male_counts, labels=['Male', 'Female'], autopct='%1.1f%%', startangle=90, colors=['#4CAF50', '#FFC107'])
    axes[0].set_title('Smartphone Usage (Male vs Female) - 2019')

    # Pie chart for Male vs Female in 2020
    axes[1].pie(female_counts, labels=['Male', 'Female'], autopct='%1.1f%%', startangle=90, colors=['#4CAF50', '#FFC107'])
    axes[1].set_title('Smartphone Usage (Male vs Female) - 2020')

    # Display the pie charts
    plt.tight_layout()
    return plt

def generate_report(output_file='report.docx'):
    # Create a new Word document
    doc = Document()

    # Chart 1: Temperature Line Chart
    fig, ax = plt.subplots()
    plot_temperature_data('temperature_data.csv', 'Date',
                         ['New York_temperature', 'London_temperature', 'New Delhi_temperature', 'Canberra_temperature'],
                         ['New York', 'London', 'New Delhi', 'Canberra'])
    plt.title('Avg Monthly Temperature of Popular 4 Cities over Last 2 Years')
    plt.savefig('temperature_chart.png')
    plt.close(fig)
    doc.add_picture('temperature_chart.png', width=Inches(5.5))

    # Chart 2: Stacked Column Chart
    fig, ax = plt.subplots()
    plot_stacked_column_chart('netflix_titles.csv')
    plt.title('Number of Movies and TV Shows Released Each Year (Starting from 2000)')
    plt.savefig('stacked_column_chart.png')
    plt.close(fig)
    doc.add_picture('stacked_column_chart.png', width=Inches(5.5))

    # Chart 3: Pie Charts for Male vs. Female Smartphone Usage
    fig, axes = plt.subplots(1, 2, figsize=(10, 5))
    plot_gender_usage_pie_chart('SmartphoneUsageDataset.csv')
    plt.suptitle('Smartphone Usage (Male vs Female) - 2019 and 2020')
    plt.savefig('gender_pie_charts.png')
    plt.close(fig)
    doc.add_picture('gender_pie_charts.png', width=Inches(5.5))

    # Save the Word document
    doc.save(output_file)

generate_report('report.docx')
